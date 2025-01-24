import os
from openai import OpenAI
from docx import Document
import comtypes.client

client = OpenAI(
    api_key=os.environ.get("OPENAI_API_KEY"),
)

def summarize_text(text):
    try:
        chat_completion = client.chat.completions.create(
            messages=[
                {
                    "role": "system",
                    "content": "You are a helpful assistant."
                },
                {
                    "role": "user",
                    "content": f"""
                    Please read the following text and provide a summary in a few sentences.
                    Focus on the main idea, key points, and emotional tone, ensuring the summary is clear and informative. Write the summary in the language of the original text.

                    Text: '{text}'
                    """
                }
            ],
            model="gpt-3.5-turbo",
        )
        response_dict = chat_completion.model_dump()
        response_message = response_dict["choices"][0]["message"]["content"]
        print(f"Summary for the text: {response_message.strip()}")
        return response_message.strip()
    except Exception as e:
        print(f"Error analyzing text: {text}. Error: {e}")
        return "unknown"
    
def create_word_and_pdf_file(text,folder_path=None ,word_filename="summary.docx",pdf_filename="summary.pdf"):
    
    if folder_path is None:
        folder_path = os.path.join(os.environ["USERPROFILE"], "software-engineering-project", "Database")
    
    os.makedirs(folder_path, exist_ok=True)

    word_filepath = os.path.join(folder_path, word_filename)
    pdf_filepath = os.path.join(folder_path, pdf_filename)

    doc = Document()
    doc.add_heading('Summary', 0)
    doc.add_paragraph(text)
    doc.save(word_filepath)
    print(f"Word file saved: {word_filepath}")
    
    # PDF conversion
    word = comtypes.client.CreateObject('Word.Application')
    try:
        word_doc = word.Documents.Open(word_filepath)
        word_doc.SaveAs(pdf_filepath, FileFormat=17)
        word_doc.Close()
        print(f"PDF file saved: {pdf_filepath}")
    except Exception as e:
        print(f"Error while converting to PDF: {e}")
    finally:
        word.Quit()

def summarize_from_txt_file(input_txt_file, folder_path=None, word_filename="summary.docx", pdf_filename="summary.pdf"):
    try:
        with open(input_txt_file, "r", encoding="utf-8") as file:
            text = file.read()
        
        print(f"Text from file '{input_txt_file}' successfully read.")
        
        summary = summarize_text(text)
        
        create_word_and_pdf_file(summary, folder_path, word_filename, pdf_filename)
    except Exception as e:
        print(f"Error while processing the file '{input_txt_file}': {e}")

input_file = "example.txt"

summarize_from_txt_file(input_file)

